"""기능명세서 양식 docx에 마중(Majung) 내용을 채워 제출본을 생성.

문체 기준: 스킬 fin-biz-tone (금융권 비즈니스 한국어, 은어 순화, em dash 금지,
완전한 구 단위 개조식). 원본 양식의 틀은 보존하고 새 파일로 출력.

실행:  cd majung/docs && python build_docx.py
"""
from __future__ import annotations

from pathlib import Path

import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parent.parent.parent  # JB_AI Challenge/
TEMPLATE = ROOT / "[데이콘] JB금융그룹 Fin AI Challenge 기능명세서 양식.docx"
OUTPUT = ROOT / "[제출]기능명세서_마중.docx"
CAPTURES = Path(__file__).resolve().parent / "captures"

FONT = "맑은 고딕"  # Windows 기본 폰트. Pretendard 미설치 환경의 대체 렌더 방지


def _set_font(run, size: float) -> None:
    run.font.size = Pt(size)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:eastAsia", "w:hAnsi"):
        rfonts.set(qn(attr), FONT)


def fill_cell(cell, lines, size: float = 9, bold: bool = False) -> None:
    """셀 내용을 lines(문단 리스트)로 교체. 셀 테두리와 음영은 보존.

    문자열 안의 개행(\\n)은 별도 문단으로 분리해 Word에서 줄바꿈이 누락되지 않게 처리.
    """
    if isinstance(lines, str):
        lines = [lines]
    expanded: list[str] = []
    for ln in lines:
        expanded.extend(ln.split("\n"))
    # 기존 문단 정리 (첫 문단만 남기고 제거)
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    p0 = cell.paragraphs[0]
    for r in list(p0.runs):
        r._element.getparent().remove(r._element)
    first = True
    for ln in expanded:
        p = p0 if first else cell.add_paragraph()
        run = p.add_run(ln)
        run.bold = bold
        _set_font(run, size)
        first = False


def append_images_row(cell, paths, width_each: float, caption: str = "") -> bool:
    """캡처 PNG들을 셀 끝에 한 문단으로 나란히 삽입. 전부 존재할 때만 True."""
    if not all(p.exists() for p in paths):
        return False
    p = cell.add_paragraph()
    for i, img in enumerate(paths):
        run = p.add_run("  " if i else "")
        _set_font(run, 9)
        p.add_run().add_picture(str(img), width=Inches(width_each))
    if caption:
        cp = cell.add_paragraph()
        run = cp.add_run(caption)
        _set_font(run, 8)
    return True


def main() -> None:
    d = docx.Document(str(TEMPLATE))
    t = d.tables

    # ── 머리표 (TABLE 0): 팀명, 주제 구분, 팀원, 작성일 ──
    for row in t[0].rows:
        cells = row.cells
        for i in range(len(cells) - 1):
            label = cells[i].text.strip()
            if label == "팀명":
                fill_cell(cells[i + 1], "마중", size=9)
            elif label.startswith("팀원 정보"):
                fill_cell(cells[i + 1], "문기훈 (1인 참가)", size=9)
            elif label == "주제 구분":
                fill_cell(cells[i + 1], "자유주제", size=9)
            elif label == "작성일":
                fill_cell(cells[i + 1], "2026.06.11", size=9)

    # ── §1 서비스 개요 (TABLE 1, KV 5행) ──
    s1 = {
        "서비스명": ["마중 (Majung)"],
        "서비스 한줄 소개": [
            "외국인 근로자의 신용 사다리 구축을 위한 맞춤형 위임(Mandate) 뱅킹 에이전트",
            "제1원칙: LLM은 의도 파싱만 수행하며, 자금 이체 실행 권한은 부여하지 않음(결정적 코드 기반 3중 게이트 통제)",
        ],
        "개발 목표": [
            "· 대상: 언어 장벽 및 금융 이력 부족으로 고금리 사채(연 20~40%)에 의존하는 외국인 근로자",
            "· 1단계: 급여 입금 트리거 및 사용자 지정 환율 조건에 기반한 자동 송금 파이프라인 구축",
            "· 2단계: 사채 이용 고객을 전북은행 대환대출(연 13.59%)로 유도하여 우량 고객화 및 신용 형성 지원",
            "· 포지셔닝: 기존 브라보코리아 앱 내장형 대환 전환 엔진(Conversion Engine)",
            "· JB 전략 연계: JB금융그룹의 AI 전환(AX) 흐름 및 외국인·전략 여신 확대 방침(그룹 외국인 여신 잔액 1조 돌파, 연말 1.3~1.5조 목표)의 직접 실행 과제",
        ],
        "타겟 사용자": [
            "· 국내 체류 외국인 273만 명 중 금융 소외 계층인 E-9(비전문취업) 근로자 30만 명 대상 (출처: 법무부, 2024)",
            "· 핵심 페르소나: 한국어에 미숙하고 고금리 사채를 이용 중인 베트남 국적 E-9 근로자",
            "· 타겟 모수: 브라보코리아 휴면 고객 20만 명 및 전북은행 기존 외국인 고객 24만 명",
        ],
        "기대 효과": [
            "[기대효과 4축 분류]",
            "· 고객가치: 대면 모집인 개입 없이 대환 가심사 및 신청을 지원하여 1인당 연간 약 250만 원 이자 비용 절감, 모국어 약관 렌더링으로 금융 접근성 향상",
            "· 업무효율: 위임 기반 자동 송금 파이프라인으로 주말 창구 수작업 제거, 대면 모집인 비용 제거로 고객획득비용(CAC) 최소화",
            "· 리스크감소: AML 자동 탐지(스코어 40·70 이중 컷)와 모국어 위임장 재확인으로 불완전판매 및 STR 누락 위험 제거",
            "· 수익기회: 기본 전환율 7% 달성 시 예상 대환 잔액 693억 원, 연 이자수익 66.5억 원 창출 (자체 추정치). 고객 생애 가치(LTV) 및 대환 전환율 증대",
        ],
    }
    for row in t[1].rows:
        label = row.cells[0].text.strip()
        if label in s1:
            fill_cell(row.cells[1], s1[label])

    # ── §2 시스템 구성도 (TABLE 2, 1x1) ──
    arch = [
        "[전체 아키텍처 흐름]",
        "사용자 발화 및 급여 입금 이벤트 발생 → LLM 기반 자연어 의도 파싱(구조화된 Intent 추출) → "
        "3중 보안 게이트(결정적 코드) 검증 → 송금 실행 엔진 가동 → 결과 검증 및 모국어 알림",
        "",
        "[핵심 통제 모듈: 3중 게이트 (All PASS 시에만 실행)]",
        "· Gate A (위임장 검증): 전자서명 무결성, 유효기간, 사용자 철회 여부, 위임 범위 초과 여부 검증",
        "· Gate B (조건 및 한도 검증): 목표 환율 도달 여부(7일 평균 대비 1% 이상) 및 1회/월간 누적 송금 한도 초과 여부 확인",
        "· Gate C (수취인 검증 및 AML 이상거래 탐지): 수취인 화이트리스트 및 블랙리스트 대조, 의심거래 스코어링(신규/고액, 분할 송금 패턴 등) → 보류, STR 큐 이관, 혹은 차단 판정",
        "",
        "[주요 컴포넌트]",
        "· 프론트엔드: React(Vite) 기반 챗봇 UI 및 관리자 대시보드",
        "· 백엔드/오케스트레이터: FastAPI 기반 MCP 호스트 (게이트웨이 역할 및 제어 로직 강제화)",
        "· AI 파이프라인: 상용 LLM(Claude 등)을 모델 추상화 계층으로 연동하도록 설계하여 자연어 의도 파싱과 모국어 번역만 담당. "
        "계좌·금액·수취인 실명은 전송하지 않고 비식별 의도 텍스트만 전달하여 개인정보 국외 이전에 해당하지 않도록 설계. "
        "RAG(벡터DB·임베딩) 미적용, 위임 이력·환율·AML 룰 판정은 모두 SQLite 기반 결정적 코드로 처리하여 환각 경로를 차단. "
        "예선 PoC 데모는 동일 인터페이스의 룰 기반 폴백으로 구현하여 API 키 없이도 전 시나리오 재현 가능. "
        "상용화 시 금융 클라우드 안전성 평가를 거친 국내 리전 또는 행내 sLLM으로 교체 가능(특정 모델 종속 없음).",
        "· 뱅킹 코어(Mock): SQLite 활용 (Mandate Service, 가상 계좌 원장 등 13개 테이블 구성)",
        "",
        "[시스템 설계 경계 및 제약 사항]",
        "① 자금 이체 등 핵심 금융 실행 권한은 LLM에게 부여하지 않음.",
        "② 대출 최종 승인은 전북은행 고유 심사엔진으로 이관(에이전트는 가심사만 수행).",
        "③ 외부 LLM 호출 시 민감 개인정보(PII)는 내부 망에서 토큰화하여 전송.",
        "④ 단일 금융 사고 발생 시 최대 손실액은 고객이 설정한 위임 한도 내로 캡핑(Capping).",
    ]
    arch_img = CAPTURES / "architecture.png"
    if not arch_img.exists():
        arch += ["", "※ 구성도 그림: 별첨(docs/diagrams.md) 캡처 삽입 예정"]
    cell_arch = t[2].rows[0].cells[0]
    fill_cell(cell_arch, arch, size=9)
    append_images_row(cell_arch, [arch_img], 6.2,
                      caption="그림 1. 시스템 구성도: LLM의 실행 권한 배제 및 결정적 코드 기반 3중 게이트 아키텍처")

    # ── §3 핵심 기능 명세 (TABLE 3, 양식 5열 유지) ──
    feats = [
        ("위임장 자동 생성 및 서명",
         "모국어 발화를 분석하여 위임장 JSON 데이터로 구조화. 건별 통지와 철회권 등 전자금융거래법 거래지시 요건을 내장. "
         "모국어 재확인 절차 누락 시 발급을 원천 차단함. [책임 경계] 위임 전문의 무결성을 해시값으로 증명. "
         "(예선 PoC) SHA-256 해시 기반 위변조 탐지. (상용화) 전자서명법 적격 수단(공인전자서명·FIDO2)으로 대체 및 금융위 유권해석 신청 예정.",
         "입력: 모국어 음성/텍스트\n출력: 위임장 JSON 데이터", "Pydantic 스키마 검증, SHA-256 무결성 해시", "구현 완료 (PoC)"),
        ("급여 입금 트리거",
         "사용자 계좌의 급여 입금 패턴을 감지하여 위임 송금 프로세스를 자동으로 시작함. 조건 미충족 시 프로세스를 대기 상태로 유지함.",
         "입력: user_id, 입금액, 연속 수령 개월\n출력: 프로세스 트리거 신호", "이벤트 기반 스트림 감지 (모의 코어)", "구현 완료"),
        ("위임장 유효성 검증 (Gate A)",
         "위임장의 전자서명, 만료일, 철회 상태 등을 내부 결정적 코드(Rule)로 검증. 위반 사항 발견 시 프로세스를 거절 또는 보류함. "
         "[책임 경계] 검증 로직 오류에 대한 책임은 시스템에 있음.",
         "입력: mandate_id, intent\n출력: 유효성 여부(Boolean), 위반 사유 배열", "결정적 규칙 엔진 (Rule Engine)", "구현 완료"),
        ("환율 조건 및 한도 검증 (Gate B)",
         "현재 환율이 7일 이동평균 대비 1% 이상 유리한지 계산. 1회 및 월간 송금 한도를 초과하는지 검증하며, "
         "환율 데이터 지연(60초 초과) 시 거래를 일시 정지함. [책임 경계] 송금 한도는 사용자가 직접 설정함.",
         "입력: 통화 쌍, amount\n출력: 환율 스프레드, 한도 준수 여부", "이동평균 산출 및 Rule 검증", "구현 완료 (데모 환경 +1.82% 적용)"),
        ("수취인 검증 및 이상거래 스코어링 (Gate C)",
         "화이트리스트 통과 시 감점, 신규 수취인 및 고액 송금, 분할 송금(Structuring) 감지 시 가산점을 부여하여 AML 스코어를 산출함. "
         "스코어 70점 이상 시 STR 후보 큐에 적재하고 보류 처리. [책임 경계] 탐지 및 차단은 시스템이, 최종 STR 보고는 준법감시인이 판단함.",
         "입력: 수취인 정보, amount, 송금 시간\n출력: AML Score, 위험 플래그, 차단 여부", "AML 룰 기반 스코어링 알고리즘", "구현 완료 (데모 환경 Score 75 산출)"),
        ("최종 송금 실행",
         "3중 게이트를 모두 통과(PASS)한 거래에 한해 모의 레일(한패스)을 통해 송금을 실행하고 영수증을 발급함. "
         "단 하나의 게이트라도 실패 시 즉시 중단함.",
         "입력: mandate_id, bnf, amount\n출력: 트랜잭션 ID, 처리 상태", "오케스트레이터 API 강제 호출", "구현 완료"),
        ("건별 모국어 실시간 통지",
         "송금 실행, 보류, 차단 등 모든 트랜잭션 결과를 사용자의 모국어로 실시간 통지하며, 화면 내 '위임 철회' 버튼을 상시 노출함.",
         "입력: user_id, 언어 코드, 메시지\n출력: 알림 발송 결과(Boolean)", "다국어 렌더링 및 푸시 알림", "구현 완료"),
        ("위임 즉시 철회",
         "사용자의 철회 요청 발생 시 해당 위임장을 즉시 무효(Revoked) 처리하며, 진행 중인 모든 트리거를 강제 중단함. "
         "[책임 경계] 철회 권한은 전적으로 고객에게 보장됨.",
         "입력: mandate_id\n출력: 상태 전이 결과(Revoked)", "데이터베이스 상태 전이 (State Transition)", "구현 완료"),
        ("대환대출 가심사 모듈 (핵심 비즈니스)",
         "비자 만료일, DSR(0.4 이하), 다중채무 여부 등 적격 필터를 거쳐 대환 시 예상 절약액을 산출함. "
         "금소법 설명의무 준수를 위해 이자 절약액과 향후 월 상환액, 연체 시 불이익을 동일한 비중으로 안내함. [책임 경계] 본 모듈은 가심사만 수행함.",
         "입력: user_id, 금융 마이데이터\n출력: 적격 여부, 이자 절약액, 예상 한도/금리", "여신 가심사 결정적 룰, 원리금 균등상환 계산", "구현 완료 (데모 기준 첫 해 절감액 2,461,500원, 약 250만 원)"),
        ("전북은행 심사엔진 회부 IF",
         "대환 가심사를 통과하고 고객이 동의한 건에 대해 전북은행 본 심사엔진(모의)으로 데이터를 이관하고 접수 번호를 발급함. "
         "[책임 경계] 최종 대출 승인 및 거절은 전북은행의 배타적 권한임.",
         "입력: user_id, 대환 신청 데이터\n출력: 심사 접수 번호", "모의 외부 API 인터페이스 연동", "구현 완료"),
        ("감사 로그 기록 (Audit Trail)",
         "모든 게이트의 판정 결과, 사유 코드, 페이로드를 조작 불가능한 형태로 적재하여 시스템의 설명 가능성(Explainability)을 확보하고 "
         "분쟁 시 방어 근거로 활용함.",
         "입력: 시스템 이벤트 데이터\n출력: 고유 로그 ID", "불변성(Immutable) 로깅 시스템", "구현 완료"),
        ("STR(의심거래보고) 후보 관리 대시보드",
         "Gate C에서 이상 징후로 탐지된 거래(Score 70 이상)를 내부 관리자 대시보드 대기열(Queue)에 자동 적재하여 모니터링을 지원함.",
         "입력: 트랜잭션 내역, 플래그\n출력: 대기열 적재 완료 상태", "관리자 큐(Queue) 매니지먼트", "구현 완료"),
    ]
    tbl = t[3]
    need = len(feats) - (len(tbl.rows) - 1)
    for _ in range(max(0, need)):
        tbl.add_row()
    for row, feat in zip(tbl.rows[1:], feats):
        for cell, val in zip(row.cells, feat):
            fill_cell(cell, val, size=8.5)

    # ── §4 주요 기능 흐름도 (TABLE 4, 1x1) ──
    flow = [
        "[As-Is 고객 여정 vs To-Be (마중 개입 단계)]",
        "· 입국·계좌 개설: (As-Is) 사업주 지정 계좌로 선택권 없음 → (To-Be) 브라보코리아 앱에서 모국어 위임장 설정 [마중 개입]",
        "· 급여·송금: (As-Is) 주말 영업점 1~2시간 대기 후 직접 송금 → (To-Be) 조건 충족 시 백그라운드 자동 송금 및 즉시 모국어 통지 [마중 개입]",
        "· 신용 형성: (As-Is) 사채 이자 누적, 제도권 거절, 사채 재유입 악순환 → (To-Be) 대환 가심사 자동 제안, 전북은행 회부, 신용 이력 형성 [마중 개입]",
        "",
        "[핵심 데모 시나리오 (End-to-End)]",
        "1. 위임 및 자동 송금: 모국어 발화(\"급여 입금 시 환율 유리할 때 200만 원 송금\") → 위임장 내용 재확인 및 서명 → "
        "급여 감지 → 환율 조건(+1.82%) 도달 시 자동 송금 실행 및 통지",
        "2. 이상 거래 보류: 심야 시간대 신규 수취인 앞 480만 원 송금 시도 → AML 스코어 75점 산출 → "
        "시스템 자동 보류 및 모국어 확인 절차 진행 → STR 대기열 적재",
        "3. 보이스피싱 원천 차단: \"OTP를 보내면 대출 가능\" 등 사기 패턴 감지 → 블랙리스트 매칭에 의한 즉각 거래 차단 및 경고 메시지 발송",
        "4. 사채 대환 유도: 기존 사채 이율 분석 → 대환 시 연 약 250만 원 절감 효과 안내 → 예상 상환액 및 연체 리스크를 동일 비중으로 고지 → "
        "전북은행 심사엔진 회부",
        "5. (확장성 테스트) 표준 MCP 인터페이스 적합성 시연. 외부 AI 플랫폼 연동 가능성을 보이되, 실제 연동 여부와 데이터 노출 범위는 전북은행 정책으로 통제(기본 비활성).",
        "",
        "[AI Agent 루프 아키텍처]",
        "· 매크로 루프: 데이터 수집 → 상태 판단(Gate A/B/C) → 액션 수행(실행/보류/알림) → 결과 검증 및 위임 한도 재조정",
        "· 마이크로 파이프라인 (평가 기준 6단계): 데이터 수집 → 조건 검색 → 룰 기반 판단 → 모국어 메시지 생성 → 사용자 검증 → 후속 대환 액션",
    ]
    uc_img = CAPTURES / "usecase.png"
    if not uc_img.exists():
        flow += ["", "※ 유스케이스 다이어그램: 별첨(usecase.puml, diagrams.md) 캡처 삽입 예정"]
    cell_flow = t[4].rows[0].cells[0]
    fill_cell(cell_flow, flow, size=9)
    append_images_row(cell_flow, [uc_img], 6.2,
                      caption="그림 2. 유스케이스 다이어그램: 시스템 액터(근로자, 관리자, 전북은행 심사엔진, 제휴처) 간 상호작용")
    append_images_row(
        cell_flow,
        [CAPTURES / f for f in ("02_auto_remit.png", "03_str_hold.png", "05_refi_offer.png")],
        2.05,
        caption="그림 3. 주요 UI: ② 자동 송금(철회 보장) / ③ 이상거래 보류(Score 75) / ⑤ 대환대출 제안(투명한 정보 고지)")

    # ── §5 향후 발전 방향 (TABLE 5, 1x1) ──
    future = [
        "[단계적 고도화 로드맵: PoC → 파일럿 → 내부망 적용 → 상용화]",
        "※ 본 예선 제출본은 PoC(모의 API 기반 E2E 검증) 범위를 충족함.",
        "· 자금 이체: (PoC) 모의 레일 구동 → (파일럿) 실시간 환율 API 연동 → (내부) 코어뱅킹 샌드박스 테스트 → (상용화) 실제 송금망 연동",
        "· 대환대출: (PoC) 자체 가심사 → (파일럿) 전북은행 심사엔진 API 연동 → (내부) 실제 여신 심사 파이프라인 연결 → (상용화) 대환 및 적금 상품 전면 운영",
        "· 규제 및 약관: (PoC) 위임 3종 스키마 적용 → (파일럿) 금융위 유권해석 신청 → (내부) 책임 배상 한도 산정 → (상용화) 공식 약관 반영",
        "· AML 고도화: (PoC) 내부 룰 스코어링 → (파일럿) STR 관리 큐 운영 → (상용화) 행내 정규 AML 시스템 완전 연계",
        "· UX 확장: (현재) 한국어/베트남어 텍스트 지원 → (상용화) 17개국 다국어 및 STT(음성 인식) 적용 추진",
        "※ 망분리, 전자금융감독규정 심사, 코어뱅킹 실연동 등은 상용화 단계의 과제로 본 대회 범위에서 제외함.",
        "",
        "[단계별 완료 기준 (검증 지표)]",
        "· PoC: End-to-End 5단계 시나리오 100% 통과, pytest 20건 전원 Green (달성)",
        "· 파일럿: 베타 사용자 50명 30일 운용, 오송금 0건, 대환 상담 전환율 5% 이상",
        "· 내부망: 코어뱅킹 샌드박스 정합성 100%, AML STR 큐 실운영 30일, 보안 취약점 점검 통과",
        "· 상용화: 전자금융거래법 유권해석 획득, 행내 AML 시스템 완전 연계, 금융감독원 전자금융 심사 승인",
        "",
        "[계열사 확장 매트릭스] (●=즉시 적용 / ○=파일럿 필요 / -=해당 없음)",
        "· 전북은행: 위임 자동송금 ●, 사채 대환대출 ●, 급여연계 적금 ●, 신용이력 이전 ○ (현재 적용 대상)",
        "· 광주은행: 위임 자동송금 ○, 사채 대환대출 ○, 급여연계 적금 ○, 신용이력 이전 ○ (그룹 내 동일 모델 확장)",
        "· JB우리캐피탈: 사채 대환대출 ●, 신용이력 이전 ○ (여신 특화 채널 연계)",
        "· 한패스(제휴): 송금 레일 공유 ●, 신용이력 이전 - (송금 데이터 연동, 2대 주주)",
        "",
        "[비즈니스 임팩트 및 성과 지표]",
        "· 재무 목표: 타겟 고객 대상 전환율 7% 달성 시, 예상 대환 잔액 693억 원 및 연간 이자수익 66.5억 원(금리 스프레드 9.59%p 적용, 대손충당금·운영비 차감 전) 창출",
        "· 마케팅 효율: 플랫폼 내재화로 대면 모집인 수수료를 제거하여 신규 고객 획득 비용(CAC) 최소화",
        "· 핵심 KPI: 단기 모객 수가 아닌 타행/사채 대환 비율 및 고객 생애 가치(LTV) 상승에 집중",
    ]
    fill_cell(t[5].rows[0].cells[0], future, size=9)

    # ── §6 부록 (TABLE 6, 1x1) ──
    appendix = [
        "[A. 규제 리스크 인식 및 선제적 대응 로드맵]",
        "· 전자금융거래법: 위임형 자동 송금의 '반복적 거래 지시' 회색지대 해소를 위해 ①모국어 서명 ②건별 통지 ③즉시 철회권을 시스템에 강제 내장함. 상용화 시 유권해석 진행 예정. 마중은 독자적 송금업·여신업을 영위하지 않으며, 실제 송금은 인가받은 한패스 소액해외송금 레일, 대출 실행은 전북은행 인가 하에서만 이루어짐.",
        "· 외국환거래법: 무증빙 연 5만 달러 한도 준수를 위해 Gate B에 누적 한도 검증 로직을, Gate C에 분할 송금(Structuring) 탐지 로직을 구현함.",
        "· 신용정보법 및 프라이버시: 외부 LLM API 호출 시 개인식별정보(PII)를 토큰화하여 데이터 유출 위험을 차단함.",
        "· 특정금융정보법(AML): 의심거래(STR) 스코어링 엔진을 구축하고, 임계치 초과 건을 관리자 대기열로 자동 분리함.",
        "· 금융소비자보호법: '부당 권유 금지' 및 '설명 의무' 준수를 위해 대출 시 절약되는 이자와 예상 월 상환액 및 연체 시 불이익을 모국어로 동일한 비중으로 고지함.",
        "",
        "[B. 기술 스택 및 오픈소스 라이선스 정책]",
        "· 백엔드 인프라: Python(PSF), FastAPI(MIT), SQLite(Public Domain) 등 상업적 이용이 허가된(Permissive) 라이선스 채택.",
        "· 프론트엔드: React / Vite (MIT 라이선스).",
        "· LLM 활용: 상용 LLM(Claude 등, Anthropic)을 모델 추상화 계층으로 연동하도록 설계하며, 예선 PoC는 동일 출력 스키마의 룰 기반 폴백으로 구현함(API 키 불필요, 본선에서 무변경 치환). 상용 연동 시 입력이 모델 학습에 사용되지 않는 Enterprise 약관을 적용함. LLM에는 비식별 의도 텍스트만 전달하고 계좌·금액·실명 등 식별정보는 전송하지 않음. RAG는 미적용(판정은 SQLite 결정적 코드 전담). 상용화 시 금융 클라우드 안전성 평가를 거친 국내 리전 또는 행내 sLLM으로 전환 가능. 데모 데이터는 전부 합성으로 개인정보 및 저작권 리스크를 배제함.",
        "",
        "[C. 6대 핵심 보안 리스크 통제 방안]",
        "① LLM 환각(Hallucination): 금융 트랜잭션 실행 권한을 LLM에서 분리하고 내부 결정적 코드(Rule)에 위임함.",
        "② 책임 소재 불명확성: 해시 기반 위임 무결성 증명과 즉시 철회권으로 전금법 제9조 배상 책임을 명확히 함 (PoC는 SHA-256 자체 무결성, 상용화 시 전자서명법 적격 수단으로 교체 및 유권해석 신청).",
        "③ 개인정보 유출: 철저한 데이터 마스킹 및 감사 로그 접근 통제 적용.",
        "④ 자금세탁(AML) 악용: Gate C의 스코어링을 통한 사전 차단 및 보류 아키텍처 적용.",
        "⑤ 설명 가능성(Explainability) 부족: 판정 사유 코드를 모국어로 렌더링하고 감사 로그에 상세 페이로드를 기록함.",
        "⑥ 외부 기술 종속성: 본체는 은행 자체 UI로 구축하며, 외부 비서 연동은 MCP 표준을 통한 옵션으로만 제공함.",
        "",
        "[D. 시스템 품질 검증 (QA)]",
        "· 게이트 판정 및 신규 API 단위 테스트 20건 100% 통과 완료 (pytest 기준).",
        "· End-to-End 5단계 시나리오 검증 완료: 위임장 발급 및 서명, 조건부 자동 송금, 이상거래 보류(480만 원 시도), "
        "사기 거래 차단(OTP 요구), 대환 가심사(연 약 250만 원 절감) 및 심사엔진 회부의 정상 동작 확인.",
        "",
        "[E. 핵심 기술 산출물 (PoC 구현 완료)]",
        "· 위임장(Mandate) JSON 스키마: beneficiary_whitelist(수취인 화이트리스트), limits(1회 및 월 한도), trigger(salary_in 급여 감지), "
        "fx_condition(7일 평균 대비 1% 이상), esign(esign_hash SHA-256 및 모국어 재확인), notification(건별 모국어 통지), revocation(무조건 철회) 등 "
        "전자금융거래법 거래지시 3요건(전자서명, 건별 통지, 철회권)을 데이터 구조에 내장함.",
        "· MCP 표준 도구 13종: get_account_balance, detect_salary_deposit, get_fx_rate, validate_mandate(Gate A), check_limits(Gate B), "
        "screen_beneficiary_aml(Gate C), execute_remittance, notify_user, revoke_mandate, get_refi_offer, refer_to_jb_engine, append_audit_log, enqueue_str_candidate. "
        "이 중 execute_remittance는 3중 게이트를 모두 통과한 경우에만 오케스트레이터가 호출하며 LLM은 직접 호출할 수 없음.",
        "· SQLite 데이터 모델 13개 테이블: users, accounts, salary_events, fx_rates, mandates, beneficiaries, transactions, blacklist_patterns, loans_external, refi_products, str_queue, audit_log, notifications.",
        "· AML 스코어 산식(Gate C): 화이트리스트 통과 시 40점 감점, 위임 범위 밖 30점, 블랙리스트(신분증·비밀번호·OTP·선입금 요구) 100점 하드컷, "
        "분할 송금(Structuring) 50점, 신규 수취인 25점, 고액(300만 원 이상) 20점, 심야 플래그. "
        "라우팅은 40점 이상 보류, 70점 이상 STR 큐 적재, 블랙리스트 즉시 차단으로 분기함. (데모: 위임 범위 밖 30 + 신규 수취인 25 + 고액 20 = 75점 보류)",
    ]
    fill_cell(t[6].rows[0].cells[0], appendix, size=8.5)

    # ── §7 기능 변경이력 (TABLE 7) ──
    log = t[7]
    entry = ("2026.06.11", "전체(최초)", "대회 예선 제출용 명세서 최초 작성", "예선 제출")
    for cell, val in zip(log.rows[1].cells, entry):
        fill_cell(cell, val, size=9)
    # 남은 빈 행 정리 (텍스트만 비움)
    for row in log.rows[2:]:
        for cell in row.cells:
            fill_cell(cell, "", size=9)

    d.save(str(OUTPUT))
    print(f"[build] 생성 완료: {OUTPUT.name}")
    print(f"[build] §3 기능행 {len(feats)}개, 표 {len(t)}개 채움. 원본 양식은 보존됨.")


if __name__ == "__main__":
    main()
